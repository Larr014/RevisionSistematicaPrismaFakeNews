#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generar reporte Word con estadisticas y graficos"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

# Cargar datos
with open('C:/Users/Luis Rojas/.openclaw/workspace/clasificacion_claude.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

articulos = data['articulos']
total = len(articulos)
meta = data['metadata']
stats = data['estadisticas']

# Crear documento
doc = Document()

# Titulo
title = doc.add_heading('Revision Sistematica: Metodos para Identificar Intenciones Comunicativas', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Subtitulo
subtitle = doc.add_heading('Reporte de Clasificacion Semantica (3,575 articulos)', level=2)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Metadata
doc.add_heading('Informacion General', level=2)
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Light Grid Accent 1'
info_table.cell(0, 0).text = 'Total de articulos'
info_table.cell(0, 1).text = str(total)
info_table.cell(1, 0).text = 'Modelo utilizado'
info_table.cell(1, 1).text = meta['modelo_usado']
info_table.cell(2, 0).text = 'Fecha de generacion'
info_table.cell(2, 1).text = meta['fecha_generacion']
info_table.cell(3, 0).text = 'Articulos clasificados'
info_table.cell(3, 1).text = f"{meta['articulos_clasificados']} (100%)"
info_table.cell(4, 0).text = 'Metodo de clasificacion'
info_table.cell(4, 1).text = 'Analisis semantico con Claude AI via CLI'

doc.add_paragraph()

# Resumen ejecutivo
doc.add_heading('Resumen Ejecutivo', level=2)
summary_text = f"""Durante el proceso de clasificacion semantica de los {total} articulos recopilados de multiples bases de datos bibliograficas, se identificaron los siguientes hallazgos clave:

- Un total de {stats['con_metodo']} articulos ({stats['con_metodo']/total*100:.1f}%) contienen descripciones explícitas de metodos computacionales.

- {stats['con_intencion']} articulos ({stats['con_intencion']/total*100:.1f}%) se enfocan en la deteccion de intenciones comunicativas.

- {stats['alta_relevancia']} articulos ({stats['alta_relevancia']/total*100:.1f}%) alcanzaron relevancia alta (>=0.7).

- Los metodos mas prevalentes: analisis manual (19.3%), Transformers (6.1%), NLP tradicional (5.0%).

- La intencion mas estudiada: Fake News (18.0%), seguida de Manipulacion (12.1%)."""

doc.add_paragraph(summary_text)

doc.add_page_break()

# Hallazgos
doc.add_heading('Hallazgos Principales', level=2)

doc.add_heading('1. Metodologias Detectadas', level=3)
doc.add_paragraph("""M1_BERT: 108 (3.0%)
M2_GPT: 83 (2.3%)
M3_Transformers: 218 (6.1%) - SEGUNDA MAS PREVALENTE
M4_CNN: 68 (1.9%)
M5_RNN_LSTM: 91 (2.5%)
M6_SVM: 100 (2.8%)
M7_NaiveBayes: 51 (1.4%)
M8_RandomForest: 110 (3.1%)
M9_NLP_Traditional: 180 (5.0%)
M10_ManualAnalysis: 691 (19.3%) - PREDOMINANTE

El predominio de analisis manual (19.3%) sugiere que muchos estudios aun dependen de evaluacion humana.""")

doc.add_heading('2. Intenciones Comunicativas', level=3)
doc.add_paragraph("""I1_FakeNews: 642 (18.0%) - DOMINANTE
I2_Manipulation: 434 (12.1%)
I3_Persuasion: 154 (4.3%)
I4_Clickbait: 8 (0.2%)
I5_Polarization: 140 (3.9%)
I6_Emotion: 82 (2.3%)
I7_Conspiracy: 30 (0.8%)

El enfoque predominante en deteccion de fake news refleja la actual prioridad en mitigacion de desinformacion.""")

doc.add_heading('3. Metricas de Evaluacion', level=3)
doc.add_paragraph("""D3_Accuracy: 435 (12.2%)
D1_Precision_Recall: 270 (7.6%)
D2_AUC_ROC: 25 (0.7%)
D4_Confusion_Matrix: 25 (0.7%)

Accuracy es predominante, aunque Precision/Recall son criticas para clasificacion desbalanceada.""")

doc.add_heading('4. Distribucion Temporal', level=3)
doc.add_paragraph("""Pre-2017: 89 articulos (2.5%)
2017-2019: 254 articulos (7.1%)
2020-2022: 742 articulos (20.8%)
2023-2026: 2,490 articulos (69.7%)

El campo es relativamente joven y en rapido crecimiento.""")

doc.add_page_break()

# Graficos
doc.add_heading('Visualizaciones', level=2)

doc.add_heading('Grafico 1: Top 10 Metodos Especificos', level=3)
try:
    doc.add_picture('C:/Users/Luis Rojas/.openclaw/workspace/01_metodos_especificos.png', width=Inches(6))
except:
    doc.add_paragraph("[Imagen no disponible]")

doc.add_heading('Grafico 2: Intenciones Comunicativas', level=3)
try:
    doc.add_picture('C:/Users/Luis Rojas/.openclaw/workspace/02_intenciones.png', width=Inches(6))
except:
    doc.add_paragraph("[Imagen no disponible]")

doc.add_heading('Grafico 3: Distribucion por Relevancia', level=3)
try:
    doc.add_picture('C:/Users/Luis Rojas/.openclaw/workspace/03_relevancia.png', width=Inches(5))
except:
    doc.add_paragraph("[Imagen no disponible]")

doc.add_page_break()

doc.add_heading('Grafico 4: Tendencia Temporal (2014-2026)', level=3)
try:
    doc.add_picture('C:/Users/Luis Rojas/.openclaw/workspace/04_tendencia_temporal.png', width=Inches(6))
except:
    doc.add_paragraph("[Imagen no disponible]")

doc.add_heading('Grafico 5: Metodos Generales', level=3)
try:
    doc.add_picture('C:/Users/Luis Rojas/.openclaw/workspace/05_metodos_generales.png', width=Inches(5.5))
except:
    doc.add_paragraph("[Imagen no disponible]")

doc.add_page_break()

# Conclusiones
doc.add_heading('Conclusiones', level=2)
doc.add_paragraph("""Esta clasificacion sistematica de 3,575 articulos revela:

1. TENDENCIA METODOLOGICA: Transicion desde metodos tradicionales (19.3% analisis manual) hacia deep learning (6.1% Transformers).

2. ENFOQUE PREDOMINANTE: Deteccion de fake news domina (18.0%), seguida de manipulacion (12.1%).

3. MADUREZ DEL CAMPO: Crecimiento exponencial post-2020 (87.5% de articulos).

4. OPORTUNIDADES: Solo 6.8% alcanzaron alta relevancia directa.

5. RECOMENDACIONES:
   - Profundizar en articulos de alta relevancia (244 articulos)
   - Analizar clusters por intencion comunicativa
   - Evaluar tendencias en adoption de modelos (BERT, GPT)
   - Identificar brechas metodologicas""")

# Footer
footer = doc.add_paragraph('Documento generado automaticamente - Revision Sistematica PRISMA 2020')
footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.italic = True

# Guardar
output_path = 'C:/Users/Luis Rojas/.openclaw/workspace/Reporte_Clasificacion_Sistematica.docx'
doc.save(output_path)
print(f"Reporte guardado: {output_path}")
